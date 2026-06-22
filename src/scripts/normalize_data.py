"""
normalize_data.py — FCRAG 2.0 Data Normalization
=================================================
Converts all raw data sources into canonical formats before chunking/indexing.

Steps:
  1. DOCX → .txt          (3GPP telecom specs)
  2. fault_narratives.jsonl → normalized FCRAGDocument schema .jsonl
  3. 180 × KPI CSV logs  → kpi_narratives.jsonl  (one narrative per fault run)
  4. OAI KPM .xlsx       → oai_kpi_baseline.csv  (for IsolationForest training)

Usage:
  python scripts/normalize_data.py
  python scripts/normalize_data.py --step 1   # run only DOCX extraction
  python scripts/normalize_data.py --step 2   # run only JSONL alignment
  python scripts/normalize_data.py --step 3   # run only CSV → narratives
  python scripts/normalize_data.py --step 4   # run only XLSX → CSV
"""

import argparse
import json
import os
import re
import uuid
from pathlib import Path

import pandas as pd
from tqdm import tqdm

# ─────────────────────────────────────────────────────────────────────────────
# Paths
# ─────────────────────────────────────────────────────────────────────────────
ROOT = Path(__file__).resolve().parent.parent.parent
DATA_RAW = ROOT / "data"
DATA_OUT = ROOT / "data" / "processed"

TELECOM_DOCS_DIR = DATA_RAW / "telecom_docs"
SIMU5G_DIR = DATA_RAW / "simu5g"
OAI_DIR = DATA_RAW / "OAI_RAN_KPM_dataset-main"

OUT_3GPP = DATA_OUT / "3gpp_text"
OUT_SIMU5G = DATA_OUT / "simu5g_docs"
OUT_KPI_NARRATIVES = DATA_OUT / "kpi_narratives"

# ─────────────────────────────────────────────────────────────────────────────
# Canonical FCRAGDocument schema
# ─────────────────────────────────────────────────────────────────────────────
def make_document(
    doc_id: str,
    source_type: str,
    source_file: str,
    clause_id: str,
    text: str,
    metadata: dict,
) -> dict:
    """Return a canonical FCRAGDocument dict."""
    return {
        "doc_id": doc_id,
        "source_type": source_type,
        "source_file": source_file,
        "clause_id": clause_id,
        "text": text.strip(),
        "metadata": metadata,
    }


# ─────────────────────────────────────────────────────────────────────────────
# STEP 1 — DOCX → Plain Text
# ─────────────────────────────────────────────────────────────────────────────
def extract_docx_to_text(docx_path: Path, out_path: Path) -> int:
    """
    Extract text from a 3GPP .docx file preserving heading hierarchy.
    Saves a .txt file where section headings are marked as [HEADING: <text>].
    Returns number of paragraphs written.
    """
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(docx_path))
    lines = []
    para_count = 0

    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        text = para.text.strip()

        if not text:
            continue

        # Mark headings so chunker can use them as clause_id metadata
        if style.startswith("Heading"):
            level = re.search(r"\d+", style)
            lvl = int(level.group()) if level else 1
            prefix = "#" * lvl
            lines.append(f"\n{prefix} {text}\n")
        else:
            lines.append(text)
            para_count += 1

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
                para_count += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text("\n".join(lines), encoding="utf-8")
    return para_count


def step1_docx_to_text():
    """Step 1: Convert all .docx 3GPP specs to plain text."""
    print("\n" + "=" * 60)
    print("STEP 1 — DOCX → Plain Text")
    print("=" * 60)

    OUT_3GPP.mkdir(parents=True, exist_ok=True)

    docx_files = list(TELECOM_DOCS_DIR.rglob("*.docx"))
    if not docx_files:
        print("  [WARNING] No .docx files found in", TELECOM_DOCS_DIR)
        return

    total_paras = 0
    for docx_path in tqdm(docx_files, desc="Extracting DOCX"):
        # e.g.  TS_38_331/38331-j20.docx  →  TS_38_331.txt
        folder_name = docx_path.parent.name        # e.g. "TS_38_331"
        out_file = OUT_3GPP / f"{folder_name}.txt"

        try:
            n = extract_docx_to_text(docx_path, out_file)
            total_paras += n
            tqdm.write(f"  ✅ {docx_path.name} → {out_file.name}  ({n} paragraphs)")
        except Exception as e:
            tqdm.write(f"  ❌ {docx_path.name} failed: {e}")

    print(f"\n  Total paragraphs extracted: {total_paras}")
    print(f"  Output → {OUT_3GPP}/")


# ─────────────────────────────────────────────────────────────────────────────
# STEP 2 — fault_narratives.jsonl → Normalized FCRAGDocument schema
# ─────────────────────────────────────────────────────────────────────────────
def step2_normalize_narratives():
    """Step 2: Align fault_narratives.jsonl to the canonical FCRAGDocument schema."""
    print("\n" + "=" * 60)
    print("STEP 2 — fault_narratives.jsonl → Normalized Schema")
    print("=" * 60)

    OUT_SIMU5G.mkdir(parents=True, exist_ok=True)

    src = SIMU5G_DIR / "fault_narratives.jsonl"
    dst = OUT_SIMU5G / "fault_narratives_normalized.jsonl"

    if not src.exists():
        print(f"  [WARNING] {src} not found. Skipping.")
        return

    docs = []
    with open(src, encoding="utf-8") as f:
        for i, line in enumerate(tqdm(f, desc="Normalizing narratives")):
            line = line.strip()
            if not line:
                continue
            raw = json.loads(line)

            # Build canonical clause_id from the first clause reference
            clauses = raw.get("clauses", [])
            clause_id = clauses[0] if clauses else ""

            doc = make_document(
                doc_id=f"simu5g_narrative_{i:04d}",
                source_type="simu5g_narrative",
                source_file=src.name,
                clause_id=clause_id,
                text=raw.get("text", ""),
                metadata={
                    "fault_type": raw.get("fault_id", ""),
                    "cell_id": raw.get("cell", ""),
                    "timestamp": raw.get("timestamp", ""),
                    "clauses": clauses,
                    "action": raw.get("action", ""),
                    "anomaly_vector": raw.get("anomaly_vector", {}),
                },
            )
            docs.append(doc)

    with open(dst, "w", encoding="utf-8") as f:
        for doc in docs:
            f.write(json.dumps(doc) + "\n")

    print(f"  ✅ {len(docs)} narratives normalized → {dst}")


# ─────────────────────────────────────────────────────────────────────────────
# STEP 3 — KPI CSV Logs → Narrative JSONL
# ─────────────────────────────────────────────────────────────────────────────
# KPI columns and their readable names
KPI_META = {
    "ho_sr":         ("Handover Success Rate",   "ratio",  "high=good"),
    "rsrp":          ("RSRP Signal Power",        "dBm",    "high=good"),
    "rsrq":          ("RSRQ Signal Quality",      "dB",     "high=good"),
    "sinr":          ("SINR",                     "dB",     "high=good"),
    "throughput":    ("DL Throughput",            "Mbps",   "high=good"),
    "prb":           ("PRB Utilization",          "ratio",  "low=good"),
    "latency":       ("User Plane Latency",       "ms",     "low=good"),
    "rrc_sr":        ("RRC Setup Success Rate",   "ratio",  "high=good"),
    "rrc_reest":     ("RRC Re-establishment Rate","ratio",  "low=good"),
    "rlf":           ("Radio Link Failure Rate",  "ratio",  "low=good"),
    "prach_fail":    ("PRACH Failure Rate",       "ratio",  "low=good"),
    "connected_ues": ("Connected UEs",            "count",  "context"),
    "ul_timing_err": ("UL Timing Error",          "ratio",  "low=good"),
}

# Most informative KPIs per fault type (for narrative focus)
FAULT_KEY_KPIS = {
    "HO_FAILURE":       ["ho_sr", "rsrp", "rlf", "rrc_reest"],
    "PRB_CONGESTION":   ["prb", "throughput", "latency", "connected_ues"],
    "PRACH_CONGESTION": ["prach_fail", "rrc_sr", "connected_ues"],
    "INTERFERENCE":     ["sinr", "rsrp", "throughput", "rlf"],
    "RLF_SPIKE":        ["rlf", "sinr", "rsrp", "ho_sr"],
    "BEAM_FAILURE":     ["rsrp", "rsrq", "sinr", "throughput"],
    "PDCP_DELAY":       ["latency", "throughput", "prb"],
    "UE_MASS_DETACH":   ["connected_ues", "rrc_sr", "ho_sr"],
    "TIMING_DRIFT":     ["ul_timing_err", "sinr", "rlf"],
    "CAPACITY_OVERFLOW":["prb", "throughput", "latency", "connected_ues"],
    "NORMAL":           ["throughput", "ho_sr", "latency"],
}


def csv_to_narrative(df: pd.DataFrame, filename: str) -> str:
    """
    Convert a KPI CSV dataframe into a human-readable narrative paragraph.
    """
    fault_type = df["fault_type"].iloc[0]
    cell_id = df["cell_id"].iloc[0]
    run_id = int(df["run_id"].iloc[0])

    fault_rows = df[df["fault_active"] == 1]
    normal_rows = df[df["fault_active"] == 0]

    # If no fault period (NORMAL runs), use full df as normal
    has_fault = len(fault_rows) > 0
    baseline_rows = normal_rows if len(normal_rows) > 0 else df

    kpi_cols = [c for c in KPI_META if c in df.columns]
    key_kpis = FAULT_KEY_KPIS.get(fault_type, kpi_cols[:4])

    # Build KPI summary sentences
    kpi_lines = []
    for kpi in key_kpis:
        if kpi not in df.columns:
            continue
        label, unit, direction = KPI_META[kpi]
        baseline_mean = baseline_rows[kpi].mean()

        if has_fault:
            fault_mean = fault_rows[kpi].mean()
            delta_pct = ((fault_mean - baseline_mean) / (abs(baseline_mean) + 1e-9)) * 100
            direction_str = "↓" if fault_mean < baseline_mean else "↑"
            kpi_lines.append(
                f"{label} {direction_str} {fault_mean:.3f} {unit} "
                f"(baseline {baseline_mean:.3f}, Δ={delta_pct:+.1f}%)"
            )
        else:
            kpi_lines.append(
                f"{label}: {baseline_mean:.3f} {unit} (stable, no fault)"
            )

    kpi_summary = "; ".join(kpi_lines) if kpi_lines else "No KPI anomaly detected."

    if has_fault:
        fault_start = fault_rows["timestamp"].iloc[0]
        fault_duration = len(fault_rows) * 0.05  # 50ms per sample
        text = (
            f"KPI Operational Log — Run {run_id:04d}: {fault_type} at {cell_id}. "
            f"Fault detected at {fault_start} lasting {fault_duration:.2f}s "
            f"({len(fault_rows)} samples affected out of {len(df)} total). "
            f"Key KPI deviations during fault: {kpi_summary}. "
            f"This represents a {fault_type.replace('_', ' ').title()} event "
            f"requiring immediate RAN intervention."
        )
    else:
        text = (
            f"KPI Operational Log — Run {run_id:04d}: NORMAL operation at {cell_id}. "
            f"No fault detected across {len(df)} samples. "
            f"Baseline metrics: {kpi_summary}."
        )

    return text


def step3_csv_to_narratives():
    """Step 3: Convert 180 KPI CSV logs into narrative JSONL documents."""
    print("\n" + "=" * 60)
    print("STEP 3 — KPI CSV Logs → kpi_narratives.jsonl")
    print("=" * 60)

    OUT_KPI_NARRATIVES.mkdir(parents=True, exist_ok=True)

    kpi_logs_dir = SIMU5G_DIR / "kpi_logs"
    csv_files = sorted(kpi_logs_dir.glob("*.csv"))

    if not csv_files:
        print(f"  [WARNING] No CSV files in {kpi_logs_dir}. Skipping.")
        return

    out_file = OUT_KPI_NARRATIVES / "kpi_narratives.jsonl"
    docs = []
    errors = []

    for csv_path in tqdm(csv_files, desc="CSV → Narratives"):
        try:
            df = pd.read_csv(csv_path)

            # Validate expected columns
            required = ["fault_active", "fault_type", "cell_id", "run_id"]
            missing = [c for c in required if c not in df.columns]
            if missing:
                errors.append(f"{csv_path.name}: missing columns {missing}")
                continue

            fault_type = df["fault_type"].iloc[0]
            cell_id = df["cell_id"].iloc[0]
            run_id = int(df["run_id"].iloc[0])

            # Parse clause references from fault type (best-effort mapping)
            clause_map = {
                "HO_FAILURE":       ["TS 38.331 §5.5.4", "TS 38.321 §5.1"],
                "PRB_CONGESTION":   ["TS 38.214 §5.1.2", "TS 38.300 §16"],
                "PRACH_CONGESTION": ["TS 38.321 §5.1", "TS 38.300 §9.2"],
                "INTERFERENCE":     ["TS 38.214 §5.2", "TS 38.300 §9"],
                "RLF_SPIKE":        ["TS 38.331 §5.3.10", "TS 38.321 §5.4"],
                "BEAM_FAILURE":     ["TS 38.331 §5.17", "TS 38.214 §5.2.1"],
                "PDCP_DELAY":       ["TS 38.323 §5.2", "TS 38.300 §6.7"],
                "UE_MASS_DETACH":   ["TS 38.413 §8.6", "TS 38.401 §8.2"],
                "TIMING_DRIFT":     ["TS 38.211 §4.3", "TS 38.133 §7.1"],
                "CAPACITY_OVERFLOW":["TS 38.300 §16", "TS 38.214 §5.1"],
                "NORMAL":           [],
            }
            clauses = clause_map.get(fault_type, [])

            narrative = csv_to_narrative(df, csv_path.name)

            doc = make_document(
                doc_id=f"kpi_log_{run_id:04d}_{fault_type}_{cell_id}",
                source_type="alarm_history",
                source_file=csv_path.name,
                clause_id=clauses[0] if clauses else "",
                text=narrative,
                metadata={
                    "fault_type": fault_type,
                    "cell_id": cell_id,
                    "run_id": run_id,
                    "clauses": clauses,
                    "total_samples": len(df),
                    "fault_samples": int(df["fault_active"].sum()),
                    "normal_samples": int((df["fault_active"] == 0).sum()),
                },
            )
            docs.append(doc)

        except Exception as e:
            errors.append(f"{csv_path.name}: {e}")

    with open(out_file, "w", encoding="utf-8") as f:
        for doc in docs:
            f.write(json.dumps(doc) + "\n")

    print(f"\n  ✅ {len(docs)} KPI narratives written → {out_file}")
    if errors:
        print(f"  ⚠️  {len(errors)} files had errors:")
        for e in errors:
            print(f"      {e}")


# ─────────────────────────────────────────────────────────────────────────────
# STEP 4 — OAI XLSX → oai_kpi_baseline.csv
# ─────────────────────────────────────────────────────────────────────────────
def step4_xlsx_to_csv():
    """Step 4: Convert OAI KPM Excel dataset to CSV for anomaly detection baseline."""
    print("\n" + "=" * 60)
    print("STEP 4 — OAI .xlsx → oai_kpi_baseline.csv")
    print("=" * 60)

    xlsx_files = list(OAI_DIR.glob("*.xlsx"))
    if not xlsx_files:
        print(f"  [WARNING] No .xlsx files in {OAI_DIR}. Skipping.")
        return

    DATA_OUT.mkdir(parents=True, exist_ok=True)
    out_file = DATA_OUT / "oai_kpi_baseline.csv"

    all_dfs = []
    for xlsx_path in tqdm(xlsx_files, desc="Reading XLSX"):
        try:
            # openpyxl backend, read all sheets
            xl = pd.ExcelFile(str(xlsx_path), engine="openpyxl")
            for sheet in xl.sheet_names:
                df = xl.parse(sheet)
                df["source_file"] = xlsx_path.name
                df["sheet_name"] = sheet
                all_dfs.append(df)
                tqdm.write(
                    f"  ✅ {xlsx_path.name} [sheet: {sheet}] → "
                    f"{len(df)} rows, {len(df.columns)} cols"
                )
        except Exception as e:
            tqdm.write(f"  ❌ {xlsx_path.name} failed: {e}")

    if all_dfs:
        combined = pd.concat(all_dfs, ignore_index=True)
        combined.to_csv(out_file, index=False)
        print(f"\n  ✅ Combined {len(combined)} rows → {out_file}")
        print(f"  Columns: {list(combined.columns)}")
    else:
        print("  [ERROR] No data extracted from XLSX files.")


# ─────────────────────────────────────────────────────────────────────────────
# Validation — spot check outputs
# ─────────────────────────────────────────────────────────────────────────────
def validate_outputs():
    """Quick spot-check on all outputs."""
    print("\n" + "=" * 60)
    print("VALIDATION — Spot Check")
    print("=" * 60)

    checks = {
        "3GPP text files":     (OUT_3GPP, "*.txt"),
        "Simu5G narratives":   (OUT_SIMU5G, "*.jsonl"),
        "KPI narratives":      (OUT_KPI_NARRATIVES, "*.jsonl"),
        "OAI baseline CSV":    (DATA_OUT, "oai_kpi_baseline.csv"),
    }

    all_ok = True
    for label, (directory, pattern) in checks.items():
        files = list(directory.glob(pattern)) if directory.exists() else []
        if files:
            sample = files[0]
            size_kb = sample.stat().st_size / 1024
            print(f"  ✅ {label}: {len(files)} file(s), sample={sample.name} ({size_kb:.1f} KB)")

            # For JSONL: verify first line is valid JSON with required fields
            if sample.suffix == ".jsonl":
                with open(sample, encoding="utf-8") as f:
                    first_line = f.readline().strip()
                if first_line:
                    try:
                        doc = json.loads(first_line)
                        required_keys = {"doc_id", "source_type", "text", "metadata"}
                        missing = required_keys - set(doc.keys())
                        if missing:
                            print(f"     ⚠️  Missing schema keys: {missing}")
                            all_ok = False
                        else:
                            print(f"     Schema OK | text preview: {doc['text'][:80]}...")
                    except json.JSONDecodeError as e:
                        print(f"     ❌ Invalid JSON: {e}")
                        all_ok = False
        else:
            print(f"  ❌ {label}: NOT FOUND at {directory}")
            all_ok = False

    print()
    if all_ok:
        print("  🎉 All outputs validated. Ready for Phase 1 chunking.")
    else:
        print("  ⚠️  Some outputs missing or malformed. Check errors above.")


# ─────────────────────────────────────────────────────────────────────────────
# Entry Point
# ─────────────────────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(
        description="FCRAG 2.0 — Data Normalization Script"
    )
    parser.add_argument(
        "--step",
        type=int,
        choices=[1, 2, 3, 4],
        help="Run only a specific step (1=DOCX, 2=JSONL, 3=CSV, 4=XLSX). "
             "Default: run all steps.",
    )
    parser.add_argument(
        "--validate-only",
        action="store_true",
        help="Only run validation check on existing outputs.",
    )
    args = parser.parse_args()

    print("\n🚀 FCRAG 2.0 — Data Normalization Pipeline")
    print(f"   Root: {ROOT}")
    print(f"   Output: {DATA_OUT}")

    if args.validate_only:
        validate_outputs()
        return

    if args.step is None or args.step == 1:
        step1_docx_to_text()

    if args.step is None or args.step == 2:
        step2_normalize_narratives()

    if args.step is None or args.step == 3:
        step3_csv_to_narratives()

    if args.step is None or args.step == 4:
        step4_xlsx_to_csv()

    validate_outputs()

    print("\n✅ Normalization complete.")
    print(f"   Next step → Phase 1: run scripts/ingest_all.py")


if __name__ == "__main__":
    main()
